from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(r"E:\wittmann_interview_ai")
OUT = ROOT / "AI-Based Smart Interview System Final Report.docx"
ASSET_DIR = ROOT / "outputs" / "final_report_assets"
ASSET_DIR.mkdir(parents=True, exist_ok=True)


def load_font(name="arial.ttf", size=24):
    path = Path(r"C:\Windows\Fonts") / name
    if path.exists():
        return ImageFont.truetype(str(path), size)
    return ImageFont.load_default()


def make_workflow():
    title_font = load_font(size=28)
    text_font = load_font(size=20)
    img = Image.new("RGB", (1500, 520), "white")
    draw = ImageDraw.Draw(img)
    title = "Workflow Diagram of AI-Based Smart Interview System"
    bbox = draw.textbbox((0, 0), title, font=title_font)
    draw.text((750 - (bbox[2] - bbox[0]) / 2, 55), title, fill=(0, 0, 0), font=title_font)
    boxes = [
        ("Candidate\nRegistration", 40),
        ("OTP\nVerification", 255),
        ("Role\nSelection", 470),
        ("Question\nAssignment", 685),
        ("AI Evaluation\nNLP + OpenCV", 900),
        ("PostgreSQL\nStorage", 1115),
        ("PDF Report\nAdmin Review", 1285),
    ]
    for text, x in boxes:
        draw.rounded_rectangle((x, 165, x + 170, 270), radius=18, fill=(226, 240, 250), outline=(31, 78, 121), width=3)
        for index, line in enumerate(text.split("\n")):
            bbox = draw.textbbox((0, 0), line, font=text_font)
            draw.text((x + 85 - (bbox[2] - bbox[0]) / 2, 190 + index * 30), line, fill=(0, 0, 0), font=text_font)
    for _, x in boxes[:-1]:
        draw.line((x + 170, 217, x + 212, 217), fill=(210, 40, 55), width=4)
        draw.polygon([(x + 212, 217), (x + 198, 209), (x + 198, 225)], fill=(210, 40, 55))
    out = ASSET_DIR / "workflow.png"
    img.save(out)
    return out


def make_architecture():
    title_font = load_font(size=28)
    text_font = load_font(size=20)
    img = Image.new("RGB", (1500, 720), "white")
    draw = ImageDraw.Draw(img)
    title = "System Architecture"
    bbox = draw.textbbox((0, 0), title, font=title_font)
    draw.text((750 - (bbox[2] - bbox[0]) / 2, 35), title, fill=(0, 0, 0), font=title_font)
    layers = [
        ("Presentation Layer", "HTML, CSS, JavaScript, Camera, Speech Input"),
        ("Application Layer", "Python Flask Routes and APIs"),
        ("AI Evaluation Layer", "Keyword Matching, TF-IDF, Communication Score, OpenCV"),
        ("Data and Report Layer", "PostgreSQL, Interview Answers, PDF Reports"),
    ]
    colors = [(231, 244, 255), (242, 248, 235), (255, 244, 230), (245, 238, 255)]
    y = 120
    for index, (heading, body) in enumerate(layers):
        draw.rounded_rectangle((180, y, 1320, y + 105), radius=18, fill=colors[index], outline=(80, 90, 100), width=3)
        draw.text((220, y + 20), heading, fill=(20, 74, 112), font=title_font)
        draw.text((220, y + 62), body, fill=(0, 0, 0), font=text_font)
        if index < len(layers) - 1:
            draw.line((750, y + 105, 750, y + 145), fill=(210, 40, 55), width=4)
            draw.polygon([(750, y + 145), (740, y + 130), (760, y + 130)], fill=(210, 40, 55))
        y += 145
    out = ASSET_DIR / "architecture.png"
    img.save(out)
    return out


def make_code_image(name, code):
    font = load_font("consola.ttf", 23)
    lines = code.splitlines()
    img = Image.new("RGB", (1450, 45 + 32 * len(lines)), (30, 34, 40))
    draw = ImageDraw.Draw(img)
    y = 22
    for line in lines:
        draw.text((26, y), line, fill=(235, 241, 247), font=font)
        y += 32
    out = ASSET_DIR / name
    img.save(out)
    return out


workflow_img = make_workflow()
architecture_img = make_architecture()
code_scoring = make_code_image(
    "code_scoring.png",
    """def score_answer(answer, question):
    kw = keyword_score(answer, question.get("keywords", []))
    sem = semantic_score(answer, question.get("ideal_answer", ""))
    comm = communication_score(answer)
    total = (kw * 0.35) + (sem * 0.45) + (comm * 0.20)
    return {"total_score": round(total * 100, 2),
            "matched_keywords": matched_keywords,
            "missing_keywords": missing_keywords}""",
)
code_bank = make_code_image(
    "code_question_bank.png",
    """def select_questions_for_role(role_slug, excluded_ids=None):
    for section, count in SECTION_COUNTS.items():
        pool = conn.execute(\"\"\"
            SELECT * FROM question_bank
            WHERE role_slug = %s AND section = %s AND active = TRUE
        \"\"\", (role_slug, section)).fetchall()
        pool.sort(key=assignment_sort_key)
        selected.extend(pool[:count])""",
)

doc = Document()
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(0.75)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
normal.font.size = Pt(12)


def set_font(run, size=12, bold=False, italic=False, color=None, name="Times New Roman"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def para(text="", align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12, bold=False, before=0, after=8, line=1.5):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    run = p.add_run(text)
    set_font(run, size=size, bold=bold)
    return p


def centered(text, size=12, bold=False, after=8):
    return para(text, WD_ALIGN_PARAGRAPH.CENTER, size, bold, 0, after, 1.2)


def h1(text):
    centered(text.upper(), 14, True, 12)


def h2(text):
    para(text, WD_ALIGN_PARAGRAPH.LEFT, 12, True, 12, 6, 1.2)


def bullet(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.15)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run("•  " + text)
    set_font(run)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement("w:" + edge)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "6")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "808080")
        borders.append(el)
    tbl_pr.append(borders)


def cell_text(cell, text, bold=False, size=10.5, color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    run = p.add_run(text)
    set_font(run, size, bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(rows, widths):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    add_borders(table)
    for row in table.rows:
        for index, width in enumerate(widths):
            row.cells[index].width = Inches(width)
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            cell = table.rows[i].cells[j]
            if i == 0:
                shade_cell(cell, "1F4E79")
                cell_text(cell, value, True, 10.5, RGBColor(255, 255, 255), WD_ALIGN_PARAGRAPH.CENTER)
            else:
                cell_text(cell, value, j == 0, 10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def add_fig(path, caption, width=6.4):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    centered(caption, 11, True, 8)


def page_break():
    doc.add_page_break()


def cover_page():
    centered("AI-BASED SMART INTERVIEW SYSTEM", 16, True, 18)
    centered("INT 500 – INTERNSHIP 4", 13, True, 6)
    centered("PROJECT REPORT", 13, True, 48)
    centered("Submitted by", 12, False, 18)
    centered("Mr. NARESH KUMAR B – E0322008", 13, True, 42)
    centered("In partial fulfilment for the award of the degree of", 12, False, 14)
    centered("BACHELOR OF TECHNOLOGY", 13, True, 8)
    centered("in", 12, False, 8)
    centered("COMPUTER SCIENCE AND ENGINEERING", 13, True, 8)
    centered("(Artificial Intelligence and Data Analytics)", 12, True, 34)
    centered("Sri Ramachandra Faculty of Engineering and Technology", 12, False, 4)
    centered("Sri Ramachandra Institute of Higher Education and Research, Porur, Chennai - 600116", 12, False, 40)
    centered("MAY 2026", 12, True, 0)


cover_page()
page_break()
cover_page()
page_break()

h1("Bonafide Certificate")
para('Certified that this project report “AI-BASED SMART INTERVIEW SYSTEM” is the bonafide record of work done by “Mr. NARESH KUMAR B – E0322008” who carried out the internship work under my supervision.')
para("", after=70)
sig = doc.add_table(rows=1, cols=2)
sig.alignment = WD_TABLE_ALIGNMENT.CENTER
cell_text(sig.cell(0, 0), "Signature of the Mentor", True, 12, align=WD_ALIGN_PARAGRAPH.CENTER)
cell_text(sig.cell(0, 1), "Signature of the HOD", True, 12, align=WD_ALIGN_PARAGRAPH.CENTER)
para("", after=30)
para("Dr. Purushothaman. R\nAssistant Professor,\nDepartment of Artificial Intelligence and Data Analytics\nSri Ramachandra Faculty of Engineering and Technology,\nSRIHER, Porur, Chennai – 600116", WD_ALIGN_PARAGRAPH.LEFT, 12, False, 0, 20, 1.2)
para("Evaluation Date:", WD_ALIGN_PARAGRAPH.LEFT, 12, False, 18, 18, 1.2)
para("Examiner 1:\n\n\nExaminer 2:", WD_ALIGN_PARAGRAPH.LEFT, 12, False, 0, 8, 1.4)
page_break()

h1("Acknowledgement")
for text in [
    "I take this opportunity to express my gratitude to Prof. T. Ragunathan, Dean and Prof. A. Saravanan, Vice Dean, Sri Ramachandra Faculty of Engineering and Technology, SRIHER, for providing all the facilities to complete this internship project successfully.",
    "I express my sincere gratitude to Dr. A. Sathya, Programme Coordinator, Department of Artificial Intelligence and Data Analytics, SRET, for providing the required support and academic guidance for carrying out this study.",
    "I also extend my sincere gratitude to Dr. R. Purushothaman, our Internship Coordinator, whose efforts in facilitating my internship were vital in this learning experience.",
    "I would like to express my deepest appreciation to my supervisor and project guide for extending help and encouragement throughout the project work. I am also thankful to Wittmann Battenfeld India Pvt Ltd for giving me an opportunity to work on a company-focused smart interview system.",
    "I am grateful to all faculty members, my parents, and my friends for their constant support and encouragement during the completion of this project.",
]:
    para(text)
page_break()

h1("Table of Contents")
contents = [
    ("ABSTRACT", "vii"), ("LIST OF FIGURES", "viii"),
    ("1 INTRODUCTION", "01"), ("1.1 Background of the Study", ""), ("1.2 Problem Statement", ""), ("1.3 Objectives of the Study", ""), ("1.4 Scope and Limitations", "03"), ("1.4.1 Scope", ""), ("1.4.2 Limitations", ""), ("1.5 Significance of the Study", ""),
    ("2 LITERATURE REVIEW", "06"), ("2.1 Introduction to the Literature", ""), ("2.2 Related Work", "07"), ("2.2.1 Comparison of Existing System", ""), ("2.3 Research Gaps", ""),
    ("3 PROPOSED METHODOLOGY", "15"), ("3.1 Workflow Diagram", ""), ("3.2 Proposed System/Methodology", ""), ("3.2.1 Candidate Registration", ""), ("3.2.2 OTP Verification", ""), ("3.2.3 Role Selection", ""), ("3.2.4 Question Assignment", ""), ("3.2.5 Answer Evaluation", ""), ("3.2.6 Face Analysis", ""), ("3.2.7 Report Generation", ""), ("3.2.8 Deployment", ""), ("3.3 Module Description", ""), ("3.4 Summary", ""),
    ("4 SYSTEM IMPLEMENTATION", "27"), ("4.1 Introduction", ""), ("4.2 Hardware and Software Specifications", ""), ("4.3 Tools / Libraries Used", ""), ("4.4 Implementation Details", ""), ("4.4.1 Backend Implementation", ""), ("4.4.2 Frontend Implementation", ""), ("4.5 Dataset Collection and Descriptions", ""), ("4.6 Summary", ""),
    ("5 RESULTS AND DISCUSSION", "31"), ("5.1 Experimental Results", ""), ("5.2 Performance Comparison / Benchmarks", ""), ("5.3 Analysis of Results", ""), ("5.4 Discussion of Key Findings", ""), ("5.5 Limitations of Results", ""),
    ("6 CONCLUSION AND FUTURE WORK", "33"), ("6.1 Summary of the Research", ""), ("6.2 Conclusion Drawn from the Results", ""), ("6.3 Contributions of the Study", ""), ("6.4 Future Work", ""), ("APPENDICES", "41"), ("REFERENCES", "42"),
]
for title, pg in contents:
    para(f"{title:<68}{pg}" if pg else title, WD_ALIGN_PARAGRAPH.LEFT, 12, False, 0, 1, 1.0)
page_break()

h1("Abstract")
for text in [
    "The AI-Based Smart Interview System is a web-based intelligent interview assessment platform developed to automate first-level candidate screening. Traditional interview screening requires manual effort, repeated evaluation, and more time when many candidates apply for similar roles. This project addresses the problem by providing a structured interview system that verifies candidates, assigns role-based questions, evaluates answers using Natural Language Processing, monitors candidate face presence using computer vision, stores interview data, and generates a professional report.",
    "The backend of the system is implemented using Python Flask, while the frontend is developed using HTML, CSS, and JavaScript. Candidate verification is performed through OTP-based authentication. After verification, candidates can select roles such as Manual Testing, Automation Testing, AI Tech Support, and Software Development. The system uses exact answer matching for objective questions and keyword matching, TF-IDF cosine similarity, and communication scoring for descriptive answers.",
    "OpenCV is used to perform basic camera monitoring such as face presence detection, centered-face validation, and multiple-face alerts. PostgreSQL is used as the database for the live question bank, candidate details, assigned interview papers, answers, scores, shortlist status, and report paths. The final output is a PDF interview report that supports admin review and shortlist decisions.",
    "The project demonstrates how AI, NLP, computer vision, and database-backed web applications can support a consistent, low-cost, and scalable candidate screening process.",
]:
    para(text)
para("KEYWORDS: Artificial Intelligence, Smart Interview System, NLP, TF-IDF, OpenCV, Flask, PostgreSQL, Candidate Screening.", WD_ALIGN_PARAGRAPH.LEFT, 12, True, 14, 0, 1.2)
page_break()

h1("List of Figures")
add_table([["S.NO", "TITLE", "PAGE.NO"], ["1", "Workflow Diagram", "16"], ["2", "System Architecture", "28"], ["3", "Code Snapshot - NLP Scoring", "40"], ["4", "Code Snapshot - Question Bank", "40"]], [0.8, 4.9, 1.2])
page_break()

h1("Chapter 1")
h1("Introduction")
para("Interview assessment is an important process for selecting candidates based on technical knowledge, reasoning ability, communication skill, and professional behaviour. In a manual interview process, evaluators ask questions, listen to candidate responses, prepare notes, compare performance, and finally prepare a decision. This becomes difficult when many candidates have to be screened within a short time.")
para("The AI-Based Smart Interview System is developed as an academic and internship project to support first-level candidate screening. The system is inspired by company-focused recruitment needs and uses a structured online interview workflow. It allows candidates to register, verify OTP, choose a role, attend an interview, and receive AI-based evaluation. The administrator can review generated reports and shortlist decisions.")
h2("1.1 Background of the Study")
para("Recruitment and technical screening are increasingly supported by digital systems. Many organizations conduct online assessments before direct HR or technical interviews. However, simple online forms do not evaluate descriptive answers, communication quality, or basic interview monitoring. Therefore, an intelligent system that combines a web application, NLP-based scoring, computer vision monitoring, and automated report generation is useful for improving consistency and reducing manual effort.")
h2("1.2 Problem Statement")
para("Manual first-level interview screening requires more time, depends heavily on evaluator availability, and may produce inconsistent feedback when many candidates are assessed. There is a need for a smart system that can conduct role-based interviews, evaluate candidate answers using measurable scoring logic, monitor basic candidate presence, store interview data, and generate structured reports for admin review.")
h2("1.3 Objectives of the Study")
for item in ["To develop a Flask-based smart interview platform for candidate screening.", "To provide OTP verification and role-based interview flow.", "To evaluate answers using exact matching, keyword matching, TF-IDF cosine similarity, and communication scoring.", "To monitor basic face presence and multiple-face conditions using OpenCV.", "To store candidate details, question bank, answers, scores, and reports in PostgreSQL.", "To generate automated PDF reports for shortlist decision support."]:
    bullet(item)
h2("1.4 Scope and Limitations")
h2("1.4.1 Scope")
para("The scope of the project includes candidate registration, OTP verification, role selection, role-wise question assignment, AI-based answer evaluation, camera monitoring, report generation, and admin report viewing. The system is suitable for academic demonstration and first-level screening support.")
h2("1.4.2 Limitations")
para("The system provides basic AI-based support and does not replace final human interview judgment. Face monitoring is limited to camera visibility and multiple-face alerts. The accuracy of descriptive scoring depends on the quality of ideal answers, keywords, and question-bank preparation. Advanced emotion analysis and deep semantic models are considered future enhancements.")
h2("1.5 Significance of the Study")
para("The project is significant because it demonstrates a practical use of AI and data analytics in recruitment screening. It combines NLP, OpenCV, Flask, PostgreSQL, and PDF reporting into one working application. The system reduces repeated manual work and provides structured feedback for candidates and administrators.")
page_break()

h1("Chapter 2")
h1("Literature Review")
h2("2.1 Introduction to the Literature")
para("The literature review focuses on automated interview assessment, online video interview analysis, NLP-based answer scoring, machine learning in recruitment, computer vision face detection, and fairness in AI-based hiring systems. These studies help to understand how verbal content, non-verbal signals, and structured scoring can support interview evaluation.")
h2("2.2 Related Work")
add_table([
    ["YEAR", "AUTHORS", "TECHNIQUE USED", "KEY CONTRIBUTION", "LIMITATION"],
    ["2015", "I. Naim et al.", "Multimodal ML", "Predicted interview performance using verbal and nonverbal cues.", "Requires large interview datasets."],
    ["2016", "L. Nguyen and D. Gatica-Perez", "Video Analytics", "Analyzed online video resumes for hirability impressions.", "Focuses more on video resume context."],
    ["2017", "L. Chen et al.", "Speech + Vision", "Used speech, facial, and language features for interview judgment.", "High computational complexity."],
    ["2018", "S. Muralidhar et al.", "NLP", "Showed importance of spoken content in candidate assessment.", "Limited visual monitoring."],
    ["2011", "F. Pedregosa et al.", "Scikit-learn", "Provided TF-IDF and similarity computation tools.", "Library only, not a complete interview system."],
    ["2000", "G. Bradski", "OpenCV", "Supported real-time image processing and face detection.", "Detection quality depends on camera and lighting."],
    ["2019", "H. Suen et al.", "HR Analytics Review", "Reviewed AI usage in personnel selection.", "Highlights governance and fairness concerns."],
], [0.75, 1.2, 1.35, 2.0, 1.6])
h2("2.2.1 Comparison of Existing System")
add_table([
    ["System Type", "Features", "Limitations", "Proposed Improvement"],
    ["Manual Interview", "Human judgment and direct interaction.", "Time-consuming and inconsistent for large candidate groups.", "Automated first-level screening and report generation."],
    ["Online Test Platform", "Objective question scoring.", "Limited descriptive answer evaluation and monitoring.", "NLP scoring and camera presence checking."],
    ["Video Interview Platform", "Records candidate responses.", "Often needs manual review and paid services.", "Local Flask-based system with low-cost AI scoring."],
    ["AI Recruitment Tools", "Advanced analytics.", "May require large datasets and paid APIs.", "Academic local system using TF-IDF, OpenCV, and PostgreSQL."],
], [1.6, 1.9, 2.0, 2.0])
h2("2.3 Research Gaps")
for item in ["Many existing systems focus only on objective test scores and do not evaluate descriptive answers.", "Some AI interview systems depend on paid APIs or large private datasets.", "Candidate monitoring is often either missing or too complex for small academic projects.", "There is a need for a simple integrated system that combines candidate workflow, NLP scoring, OpenCV monitoring, database storage, and automated reports."]:
    bullet(item)
page_break()

h1("Chapter 3")
h1("Proposed Methodology")
h2("3.1 Workflow Diagram")
add_fig(workflow_img, "Figure 1: Workflow Diagram", 6.7)
h2("3.2 Proposed System/Methodology")
para("The proposed methodology follows a structured pipeline beginning with candidate registration and ending with report generation. Each module is implemented separately and connected through Flask routes and APIs. Candidate information and interview results are stored in the database, while scoring logic and camera monitoring are handled by separate utility modules.")
for sub, text in [
    ("3.2.1 Candidate Registration", "The candidate enters name, email, phone number, and resume details through the web interface."),
    ("3.2.2 OTP Verification", "A six-digit OTP is generated and verified before the candidate can proceed to the interview."),
    ("3.2.3 Role Selection", "The candidate selects a role such as Manual Testing, Automation Testing, AI Tech Support, or Software Development."),
    ("3.2.4 Question Assignment", "The PostgreSQL question bank provides role-wise aptitude and programming questions. The exact assigned paper is stored for consistency."),
    ("3.2.5 Answer Evaluation", "Objective answers are evaluated using exact matching. Descriptive answers are evaluated using keyword matching, TF-IDF semantic similarity, and communication score."),
    ("3.2.6 Face Analysis", "The webcam frame is sent to the backend API and OpenCV detects face presence, centered face status, and multiple-face alerts."),
    ("3.2.7 Report Generation", "The final score, feedback, confidence index, and shortlist decision are added to a PDF report."),
    ("3.2.8 Deployment", "The application runs as a Flask web application and can be accessed locally or through a LAN address with HTTPS support."),
]:
    h2(sub)
    para(text)
h2("3.3 Module Description")
add_table([["Module", "Description"], ["Authentication Module", "Handles candidate registration and OTP verification."], ["Question Bank Module", "Stores and selects role-based aptitude and programming questions."], ["Scoring Module", "Computes exact answer, keyword, TF-IDF, communication, and total score."], ["Computer Vision Module", "Performs webcam-based face presence and multiple-face detection."], ["Report Module", "Generates candidate-wise PDF interview reports."], ["Admin Module", "Provides database and report viewing support for admin review."]], [2.1, 4.8])
h2("3.4 Summary")
para("The methodology integrates user interaction, AI evaluation, database operations, and reporting into a single workflow. This makes the system suitable for demonstrating AI-assisted first-level interview screening.")
page_break()

h1("Chapter 4")
h1("System Implementation")
h2("4.1 Introduction")
para("System implementation converts the proposed methodology into a working Flask application. The implementation includes frontend pages, backend routes, NLP scoring utilities, OpenCV frame analysis, PostgreSQL database functions, and PDF report generation.")
h2("4.2 Hardware and Software Specifications")
h2("4.2.1 Hardware Requirements")
add_table([["Component", "Minimum Requirement"], ["Processor", "Intel i3 or above"], ["RAM", "4 GB or above"], ["Storage", "500 MB free space for application and reports"], ["Camera", "Webcam for face monitoring"], ["Network", "Localhost or LAN access for demo"]], [2.3, 4.6])
h2("4.2.2 Software Requirements")
add_table([["Software", "Purpose"], ["Python", "Backend programming language"], ["Flask", "Web application framework"], ["PostgreSQL", "Database for question bank and interview records"], ["OpenCV", "Face detection and camera monitoring"], ["Scikit-learn", "TF-IDF vectorization and cosine similarity"], ["ReportLab", "PDF report generation"], ["HTML, CSS, JavaScript", "Frontend interface and camera/speech support"]], [2.3, 4.6])
h2("4.3 Tools / Libraries Used")
para("The major tools and libraries used in this project are Python, Flask, psycopg, PostgreSQL, OpenCV, scikit-learn, ReportLab, HTML, CSS, JavaScript, and browser speech recognition support. These tools were selected because they are suitable for local development and do not require paid AI APIs for the core functionality.")
h2("4.4 Implementation Details")
h2("4.4.1 Backend Implementation")
para("The backend is implemented in app.py using Flask. It contains routes for home page, OTP verification, role selection, interview page, admin database view, admin reports, question API, frame analysis API, submit API, result page, report download, and resume download. Utility files are used for scoring, database operations, question-bank selection, and report creation.")
h2("4.4.2 Frontend Implementation")
para("The frontend is implemented using HTML templates, CSS styling, and JavaScript. The interview page displays questions, captures answers, handles camera frames, updates face status, supports browser speech input, and submits the interview to the backend API.")
h2("4.5 Dataset Collection and Descriptions")
para("The dataset for this project is the role-wise question bank prepared for the smart interview system. It contains aptitude and programming questions for roles such as Manual Testing, Automation Testing, AI Tech Support, and Software Development. Each question includes section, topic, difficulty, question text, options, correct answer, keywords, marks, and active status.")
add_fig(architecture_img, "Figure 2: System Architecture", 6.5)
h2("4.6 Summary")
para("The implementation stage completed the major functional modules required for second and final review demonstration. The system currently supports candidate flow, AI scoring, camera monitoring, database storage, and automated report generation.")
page_break()

h1("Chapter 5")
h1("Results and Discussion")
h2("5.1 Experimental Results")
para("The system was tested using sample candidate entries and role-based interview flows. Candidate registration, OTP verification, role selection, question loading, answer submission, AI scoring, face monitoring, report generation, and admin report viewing were verified successfully.")
add_table([["Feature Tested", "Result"], ["Candidate Registration", "Working"], ["OTP Verification", "Working"], ["Role Selection", "Working"], ["Question Loading", "Working from PostgreSQL question bank"], ["NLP Answer Scoring", "Working using exact, keyword, TF-IDF, and communication score"], ["Face Monitoring", "Working using OpenCV face detection"], ["PDF Report Generation", "Working"], ["Admin Report View", "Working"]], [2.6, 4.3])
h2("5.2 Performance Comparison / Benchmarks")
add_table([["Evaluation Method", "Purpose", "Expected Output"], ["Exact Answer Matching", "Aptitude and programming objective answers", "Correct / Incorrect score"], ["Keyword Matching", "Technical keyword coverage", "Matched and missing keywords"], ["TF-IDF Cosine Similarity", "Semantic similarity between candidate answer and ideal answer", "Similarity score"], ["Communication Score", "Answer length and vocabulary variety", "Communication quality score"], ["OpenCV Face Detection", "Candidate presence monitoring", "Face status and alert level"]], [2.2, 2.4, 2.3])
h2("5.3 Analysis of Results")
para("The results show that the system can complete the full assessment cycle without manual calculation. Objective questions are scored quickly using exact matching, while descriptive answers receive a combined score based on keyword relevance, semantic similarity, and communication quality. The camera module adds a basic confidence and monitoring layer to the interview process.")
h2("5.4 Discussion of Key Findings")
for item in ["The integrated workflow reduces repeated manual effort in first-level screening.", "Role-wise question selection makes the interview more relevant to the candidate profile.", "The database-backed question bank makes the system easier to update and maintain.", "Automated reports provide structured information for admin review and shortlist decisions."]:
    bullet(item)
h2("5.5 Limitations of Results")
para("The current results are based on rule-based and classical NLP techniques. The system does not perform deep emotion analysis or advanced transformer-based semantic evaluation. Camera monitoring is limited by lighting, camera quality, and browser permissions. Human evaluation is still required for final hiring decisions.")
page_break()

h1("Chapter 6")
h1("Conclusion and Future Work")
h2("6.1 Summary of the Research")
para("This project developed an AI-Based Smart Interview System for automating first-level candidate screening. The system combines Flask, PostgreSQL, NLP scoring, OpenCV monitoring, and PDF report generation to provide a complete interview assessment workflow.")
h2("6.2 Conclusion Drawn from the Results")
para("The project concludes that a smart interview system can support consistent and structured candidate screening. It helps reduce manual scoring effort, improves report preparation, and provides a repeatable process for role-based interviews. The system is suitable as an academic demonstration of AI, NLP, computer vision, and data analytics in recruitment.")
h2("6.3 Contributions of the Study")
for item in ["Designed and implemented a complete web-based interview workflow.", "Integrated NLP scoring using exact answer matching, keyword matching, TF-IDF, and communication score.", "Added OpenCV-based camera monitoring for face presence and multiple-face alerts.", "Created a PostgreSQL-backed editable question bank and interview record system.", "Generated automated PDF reports for admin review."]:
    bullet(item)
h2("6.4 Future Work")
for item in ["Add HR/admin login with role-based access control.", "Improve semantic scoring using BERT or sentence-transformer models.", "Add dashboard analytics for candidate performance comparison.", "Expand the question bank with more company-specific and role-specific questions.", "Add video/audio recording support with stronger consent and privacy controls.", "Deploy the application on a secure cloud server for real-time use."]:
    bullet(item)
page_break()

h1("Appendices")
h2("Appendix 1 – Code")
add_fig(code_scoring, "Figure 3: Code Snapshot - NLP Scoring", 6.5)
add_fig(code_bank, "Figure 4: Code Snapshot - Question Bank", 6.5)
h2("Appendix 2 – Screenshots")
para("Screenshots of the candidate registration page, OTP page, role selection page, interview page, result page, admin database page, and generated report can be added here after the final demo screenshots are captured.")
page_break()

h1("References")
refs = [
    "I. Naim et al., “Automated Prediction and Analysis of Job Interview Performance,” 2015.",
    "L. Nguyen and D. Gatica-Perez, “Hirability in the Wild: Analysis of Online Video Resumes,” 2016.",
    "L. Chen et al., “Automated Video Interview Judgment on Large Corpus,” 2017.",
    "S. Muralidhar et al., “Words Worth: Verbal Content and Hirability Impressions,” 2018.",
    "F. Pedregosa et al., “Scikit-learn: Machine Learning in Python,” 2011.",
    "K. Sparck Jones, “A Statistical Interpretation of Term Specificity,” 1972.",
    "G. Bradski, “The OpenCV Library,” 2000.",
    "J. Devlin et al., “BERT: Pre-training of Deep Bidirectional Transformers,” 2019.",
    "H. Suen et al., “Artificial Intelligence in Personnel Selection,” 2019.",
    "N. Guenole et al., “Algorithmic Recruitment and Selection Systems,” 2022.",
    "Python Software Foundation, Python Documentation.",
    "Flask Documentation, Pallets Projects.",
    "PostgreSQL Global Development Group, PostgreSQL Documentation.",
    "ReportLab User Guide for PDF Generation.",
]
for index, ref in enumerate(refs, 1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.first_line_indent = Inches(-0.3)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(f"[{index}] {ref}")
    set_font(run)

for section in doc.sections:
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "1"
    r.append(t)
    fld.append(r)
    footer._p.append(fld)

doc.save(OUT)
print(OUT)
