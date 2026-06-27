from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(r"E:\wittmann_interview_ai")
OUT = ROOT / "AI-Based Smart Interview System IEEE Paper.docx"
ASSET_DIR = ROOT / "outputs" / "conference_paper" / "assets"
ASSET_DIR.mkdir(parents=True, exist_ok=True)


def font(name="times.ttf", size=24):
    path = Path(r"C:\Windows\Fonts") / name
    if path.exists():
        return ImageFont.truetype(str(path), size)
    return ImageFont.load_default()


def make_architecture():
    title_font = font("timesbd.ttf", 30)
    text_font = font("times.ttf", 24)
    img = Image.new("RGB", (1200, 720), "white")
    d = ImageDraw.Draw(img)
    title = "Architecture of the Smart Interview System"
    bbox = d.textbbox((0, 0), title, font=title_font)
    d.text((600 - (bbox[2] - bbox[0]) / 2, 35), title, fill=(0, 0, 0), font=title_font)
    layers = [
        ("Candidate Interface", "Registration, OTP, role selection, answers, camera stream"),
        ("Flask Application Layer", "Routes, sessions, APIs, validation and orchestration"),
        ("AI Evaluation Layer", "Exact match, keyword score, TF-IDF similarity, OpenCV analysis"),
        ("Data and Report Layer", "PostgreSQL records, assigned paper, PDF report, admin review"),
    ]
    y = 115
    for heading, body in layers:
        d.rounded_rectangle((110, y, 1090, y + 105), radius=14, outline=(0, 0, 0), width=3, fill=(242, 247, 251))
        d.text((145, y + 18), heading, fill=(0, 0, 0), font=title_font)
        d.text((145, y + 60), body, fill=(0, 0, 0), font=text_font)
        if y < 550:
            d.line((600, y + 105, 600, y + 140), fill=(0, 0, 0), width=3)
            d.polygon([(600, y + 140), (590, y + 125), (610, y + 125)], fill=(0, 0, 0))
        y += 140
    out = ASSET_DIR / "architecture.png"
    img.save(out)
    return out


def make_workflow():
    title_font = font("timesbd.ttf", 28)
    text_font = font("times.ttf", 22)
    img = Image.new("RGB", (1200, 430), "white")
    d = ImageDraw.Draw(img)
    boxes = [
        ("Register", 35),
        ("OTP", 200),
        ("Role", 365),
        ("Questions", 530),
        ("AI Score", 695),
        ("Database", 860),
        ("Report", 1025),
    ]
    for text, x in boxes:
        d.rounded_rectangle((x, 155, x + 135, 230), radius=12, outline=(0, 0, 0), width=2, fill=(245, 245, 245))
        bbox = d.textbbox((0, 0), text, font=text_font)
        d.text((x + 67 - (bbox[2] - bbox[0]) / 2, 178), text, fill=(0, 0, 0), font=text_font)
    for _, x in boxes[:-1]:
        d.line((x + 135, 192, x + 162, 192), fill=(0, 0, 0), width=2)
        d.polygon([(x + 162, 192), (x + 152, 186), (x + 152, 198)], fill=(0, 0, 0))
    title = "End-to-End Interview Assessment Workflow"
    bbox = d.textbbox((0, 0), title, font=title_font)
    d.text((600 - (bbox[2] - bbox[0]) / 2, 55), title, fill=(0, 0, 0), font=title_font)
    out = ASSET_DIR / "workflow.png"
    img.save(out)
    return out


architecture_img = make_architecture()
workflow_img = make_workflow()


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.75)
section.bottom_margin = Inches(1.0)
section.left_margin = Inches(0.63)
section.right_margin = Inches(0.63)
section.page_width = Inches(8.5)
section.page_height = Inches(11)


def set_columns(section, count=1, space_twips="360"):
    sect_pr = section._sectPr
    cols = sect_pr.xpath("./w:cols")
    cols = cols[0] if cols else OxmlElement("w:cols")
    cols.set(qn("w:num"), str(count))
    cols.set(qn("w:space"), space_twips)
    if not sect_pr.xpath("./w:cols"):
        sect_pr.append(cols)


set_columns(section, 1)
normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
normal.font.size = Pt(10)


def set_font(run, size=10, bold=False, italic=False, name="Times New Roman"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def p(text="", align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=10, bold=False, italic=False, before=0, after=3, line=1.0):
    para = doc.add_paragraph()
    para.alignment = align
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after = Pt(after)
    para.paragraph_format.line_spacing = line
    run = para.add_run(text)
    set_font(run, size=size, bold=bold, italic=italic)
    return para


def heading(text):
    return p(text.upper(), WD_ALIGN_PARAGRAPH.CENTER, 10, False, False, 9, 4, 1.0)


def subheading(text):
    return p(text, WD_ALIGN_PARAGRAPH.LEFT, 10, False, True, 5, 2, 1.0)


def run_para(parts, align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=3):
    para = doc.add_paragraph()
    para.alignment = align
    para.paragraph_format.space_after = Pt(after)
    para.paragraph_format.line_spacing = 1.0
    for text, bold, italic in parts:
        run = para.add_run(text)
        set_font(run, 10, bold, italic)
    return para


def borders(table):
    tbl_pr = table._tbl.tblPr
    tbl_borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement("w:" + edge)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        tbl_borders.append(el)
    tbl_pr.append(tbl_borders)


def cell_text(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    para = cell.paragraphs[0]
    para.alignment = align
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.0
    run = para.add_run(text)
    set_font(run, 8, bold)


def add_table(rows, widths):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    borders(table)
    for row in table.rows:
        for i, width in enumerate(widths):
            row.cells[i].width = Inches(width)
    for i, row in enumerate(rows):
        for j, text in enumerate(row):
            cell_text(table.rows[i].cells[j], text, i == 0, WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT)
    p("", after=1)


def add_fig(path, caption, width=3.05):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_after = Pt(1)
    para.add_run().add_picture(str(path), width=Inches(width))
    p(caption, WD_ALIGN_PARAGRAPH.CENTER, 8, False, False, 0, 5, 1.0)


# Header line and title block.
p("AI-Based Smart Interview System for Automated Candidate Screening using NLP, Computer Vision and PostgreSQL", WD_ALIGN_PARAGRAPH.CENTER, 24, False, False, 0, 12, 1.0)
p("Naresh Kumar B", WD_ALIGN_PARAGRAPH.CENTER, 10, False, False, 0, 0, 1.0)
p("Department of Computer Science and Engineering", WD_ALIGN_PARAGRAPH.CENTER, 9, False, False, 0, 0, 1.0)
p("Artificial Intelligence and Data Analytics", WD_ALIGN_PARAGRAPH.CENTER, 9, False, False, 0, 0, 1.0)
p("Sri Ramachandra Faculty of Engineering and Technology", WD_ALIGN_PARAGRAPH.CENTER, 9, False, False, 0, 0, 1.0)
p("Sri Ramachandra Institute of Higher Education and Research, Chennai, India", WD_ALIGN_PARAGRAPH.CENTER, 9, False, False, 0, 0, 1.0)
p("Chennai, Tamil Nadu, India", WD_ALIGN_PARAGRAPH.CENTER, 9, False, False, 0, 10, 1.0)

two_col = doc.add_section(WD_SECTION.CONTINUOUS)
two_col.top_margin = Inches(0.72)
two_col.bottom_margin = Inches(1.0)
two_col.left_margin = Inches(0.63)
two_col.right_margin = Inches(0.63)
two_col.page_width = Inches(8.5)
two_col.page_height = Inches(11)
set_columns(two_col, 2, "240")

run_para([
    ("Abstract — ", True, True),
    ("Recruitment teams and academic placement cells frequently conduct first-level interviews for a large number of candidates. Manual screening is time-consuming and can produce inconsistent feedback when evaluators assess similar profiles repeatedly. This paper presents an AI-Based Smart Interview System that automates the initial interview assessment workflow using a Flask web application, Natural Language Processing based answer scoring, OpenCV based camera monitoring, PostgreSQL backed question management and automated report generation. The system verifies candidates using OTP, assigns role-based aptitude and programming questions, evaluates submitted answers through exact answer matching, keyword coverage, TF-IDF cosine similarity and communication scoring, and generates a structured PDF report for administrator review. The proposed system provides a low-cost and locally deployable approach for consistent candidate screening without depending on paid external AI services.", False, False),
], after=6)
run_para([
    ("Keywords — ", True, True),
    ("Smart interview system; candidate screening; natural language processing; TF-IDF; OpenCV; Flask; PostgreSQL; automated assessment.", False, False),
], after=8)

heading("I. Introduction")
p("Interview screening is an important stage in recruitment because it helps an organization identify candidates with suitable technical knowledge, reasoning ability, communication skill and professional behaviour. In a conventional process, evaluators ask questions, listen to responses, manually compare answers and prepare feedback. This process is useful but becomes difficult when a large number of candidates have to be screened for similar entry-level roles.")
p("Online assessments reduce some manual effort, but many systems only support objective question scoring. They do not evaluate descriptive answers, do not provide structured feedback and do not store complete interview evidence for later review. Similarly, video interview platforms can record candidate responses, but they often require manual review or paid cloud-based analytics. Hence, there is a need for a simple, transparent and locally deployable smart interview system that combines web-based interview flow, answer evaluation, basic camera monitoring, database storage and report generation.")
p("The proposed AI-Based Smart Interview System addresses this requirement. It supports candidate registration, OTP verification, role selection, role-wise question assignment, NLP based scoring, OpenCV based face presence monitoring and automated PDF report generation. The system is designed for roles such as Manual Testing, Automation Testing, AI Tech Support and Software Development. It is suitable for academic demonstration, internship review and first-level screening support.")
p("The main contribution of this work is the integration of multiple simple but useful components into a single screening framework. Instead of treating interview evaluation as only an online test or only a video interview, the system combines identity verification, question assignment, answer scoring, camera status analysis, database persistence and report generation. This makes the workflow easier to audit and easier to demonstrate in an academic environment.")
p("Another contribution is the use of explainable scoring. Since first-level screening decisions should be understandable to evaluators, the system reports matched keywords, missing keywords, semantic similarity and communication score instead of producing only a hidden model output. This supports human review and helps the administrator understand why a candidate received a particular score.")

heading("II. Related Work")
p("Automated interview assessment has been studied using multiple data sources including verbal responses, facial cues, speech features and video behaviour. Naim et al. investigated automated prediction of interview performance using multimodal features. Nguyen and Gatica-Perez studied online video resumes and hirability impressions. Chen et al. explored automated judgement of video interviews using speech, facial and language characteristics. These works show that candidate assessment can benefit from computational analysis, but they often require large datasets and complex models.")
p("NLP based answer assessment has also been widely used in educational and recruitment systems. Keyword matching and TF-IDF similarity provide explainable ways to compare candidate answers with expected responses. Scikit-learn provides practical tools for vectorization and cosine similarity. OpenCV supports real-time computer vision tasks such as face detection and frame analysis. AI recruitment studies also highlight the need for fairness, transparency and human oversight. The proposed system follows an explainable, low-cost approach by using classical NLP scoring and basic face monitoring rather than opaque decision-making models.")
p("Existing recruitment tools usually fall into three categories: manual interview systems, online test portals and AI-enabled video interview platforms. Manual systems provide direct human judgement but require more time. Test portals are scalable but are usually limited to multiple-choice evaluation. Video interview platforms preserve richer candidate evidence but are often expensive and require manual review. The proposed system attempts to occupy a practical middle ground by providing a lightweight, locally deployable and explainable assessment process.")
add_table([
    ["Approach", "Strength", "Limitation"],
    ["Manual interview", "Direct human judgement", "High time and evaluator dependency"],
    ["Online test portal", "Fast objective scoring", "Limited descriptive answer analysis"],
    ["Video interview platform", "Captures candidate behaviour", "Often paid and review-heavy"],
    ["Proposed system", "Integrated NLP, CV, database and report flow", "Basic AI methods; final HR judgement required"],
], [0.88, 1.08, 1.04])

heading("III. Proposed System Framework")
p("The proposed system is organized as a layered web application. The presentation layer contains HTML, CSS and JavaScript pages for registration, OTP verification, role selection, interview answering, camera capture and result display. The application layer is implemented using Flask routes and APIs. The AI evaluation layer performs answer scoring and camera-frame analysis. The data and report layer uses PostgreSQL for persistent records and ReportLab for PDF report generation.")
add_fig(architecture_img, "Fig. 1. Architecture of the proposed smart interview system.", 3.25)
subheading("A. Candidate Verification and Role Selection")
p("The candidate first submits name, email, phone number and resume details. A one-time password is generated and verified before the candidate is allowed to continue. After successful verification, the candidate selects an interview role. The role determines the question pattern and the topics used for assessment.")
subheading("B. Question Bank and Assignment")
p("The question bank is stored in PostgreSQL. Each record contains role slug, section, topic, difficulty, question text, options, correct answer, keywords, marks and active status. When an interview starts, the system selects active questions for the selected role and records the exact assigned paper. This prevents later question-bank edits from changing completed interview records.")
subheading("C. NLP Scoring Model")
p("Objective answers are evaluated using exact answer matching. Descriptive answers are evaluated using keyword score, semantic similarity and communication score. The total descriptive score is calculated as a weighted combination of these three components:")
p("Total Score = 0.35(KS) + 0.45(TS) + 0.20(CS)", WD_ALIGN_PARAGRAPH.CENTER, 9, False, False, 3, 3, 1.0)
p("where KS denotes keyword score, TS denotes TF-IDF cosine similarity score and CS denotes communication score. This formula gives higher importance to semantic relevance while still rewarding important technical keywords and answer clarity.")
subheading("D. Camera Monitoring")
p("The browser captures webcam frames and sends encoded images to the Flask API. OpenCV converts the frame to grayscale and applies Haar cascade based face detection. The system records whether a face is detected, whether more than one face is present and whether the candidate remains centered. These signals are used as a basic confidence and monitoring indicator.")
add_fig(workflow_img, "Fig. 2. End-to-end interview assessment workflow.", 3.25)
subheading("E. Report Generation and Administrator Review")
p("After submission, the scoring results and camera summary are combined into a candidate-level output. The report contains candidate identity, selected role, total score, question-wise performance, feedback, confidence index and shortlist status. The report is stored in the project report directory and can be accessed through the administrator report page. This reduces the time needed to manually prepare interview notes after every candidate attempt.")
subheading("F. Algorithmic Flow")
add_table([
    ["Step", "Operation"],
    ["1", "Register candidate and verify OTP"],
    ["2", "Select role and fetch active questions from PostgreSQL"],
    ["3", "Store assigned paper for the interview session"],
    ["4", "Collect answers and camera frames during the test"],
    ["5", "Compute exact, keyword, TF-IDF and communication scores"],
    ["6", "Compute face presence and multiple-face summary"],
    ["7", "Generate result, shortlist decision and PDF report"],
], [0.45, 2.45])

heading("IV. Implementation")
p("The backend implementation is developed in Python using Flask. The main application file contains routes for home page, OTP verification, role selection, interview page, database view, admin reports, question API, frame analysis API, submit API, result page and report download. Utility modules are used for scoring, database operations, question-bank selection and PDF report generation.")
p("The frontend uses HTML templates, CSS and JavaScript. The interview page loads questions, records answer selections, displays camera status, supports browser speech input and submits responses through an API call. The scoring module uses scikit-learn's TF-IDF vectorizer and cosine similarity. The database module uses PostgreSQL through psycopg. Report generation is handled using ReportLab.")
p("The implementation follows a modular organization. The scoring functions are placed in a separate utility file so that the answer evaluation logic can be tested and improved independently. The database functions create and manage the tables required for roles, users, interviews, questions and candidate answers. The question-bank module is responsible for selecting the least-used active questions for a selected role. This separation improves maintainability and makes the application easier to extend.")
add_table([
    ["Layer", "Main Components", "Responsibility"],
    ["Frontend", "HTML, CSS, JavaScript", "Candidate forms, question display, camera capture and answer submission"],
    ["Backend", "Flask routes and APIs", "Session handling, OTP workflow, validation and orchestration"],
    ["AI Evaluation", "Scikit-learn, OpenCV", "NLP scoring, semantic similarity, face detection and alerts"],
    ["Database", "PostgreSQL", "Question bank, users, interviews, assigned papers and answer records"],
    ["Report", "ReportLab", "Score summary, feedback, confidence index and PDF report"],
], [0.72, 0.9, 1.28])
p("The database design gives importance to auditability. The live question bank is editable, but the exact questions assigned to a candidate are copied into interview records. Therefore, if an administrator later edits or disables a question, the completed candidate record remains unchanged. This behaviour is important in interview systems because old results must remain traceable and consistent.")
add_table([
    ["Data Entity", "Purpose"],
    ["users", "Candidate identity, contact information, OTP status and resume path"],
    ["roles", "Supported interview roles and role names"],
    ["question_bank", "Editable active questions with topic, answer, keywords and marks"],
    ["interviews", "Interview id, candidate id, selected role, status and report path"],
    ["candidate_answers", "Submitted answers, AI score, feedback and keyword details"],
], [0.95, 1.95])

heading("V. Results and Discussion")
p("The system was tested using sample candidate flows for role-based interviews. Candidate registration, OTP verification, role selection, question loading, answer submission, NLP scoring, face monitoring, database storage, PDF report generation and admin report viewing were verified. The intermediate implementation demonstrates approximately eighty percent completion of the overall project workflow.")
p("For testing, candidate attempts were created for different roles. Each test run verified whether the expected questions were loaded, whether answers were accepted by the submit API, whether the score summary was generated and whether report paths were stored correctly. Camera analysis was tested by checking normal face presence, missing-face condition and multiple-face alert conditions. These tests confirmed that the major modules can interact as one end-to-end workflow.")
add_table([
    ["Feature", "Observed Result"],
    ["OTP verification", "Candidate can proceed only after verification"],
    ["Role-wise questions", "Questions are loaded from the PostgreSQL bank"],
    ["NLP evaluation", "Exact, keyword, TF-IDF and communication scores are generated"],
    ["Camera monitoring", "Face missing and multiple-face states are detected"],
    ["Report generation", "PDF report is created for administrator review"],
], [1.15, 1.75])
p("The key advantage of the system is that it provides a consistent screening process while keeping the scoring logic explainable. PostgreSQL based question management makes the question bank editable and reusable. The generated report reduces manual documentation effort. However, the system does not replace final human interview judgement. The current computer vision module provides only basic monitoring and the NLP model uses classical similarity techniques. Advanced transformer-based semantic scoring, dashboard analytics and secure cloud deployment can improve the system further.")
subheading("A. Discussion")
p("The observed behaviour shows that a classical NLP based approach can be useful when the goal is first-level screening rather than final selection. Exact answer matching works well for objective aptitude and programming questions. Keyword matching helps verify whether the candidate has mentioned important technical terms. TF-IDF cosine similarity gives an approximate semantic comparison between the submitted answer and the expected answer. Communication score rewards answers that are neither too short nor excessively unclear.")
p("The OpenCV component is intentionally kept simple. It does not infer emotion or personality; it only provides a basic indication of whether a candidate is visible and whether more than one face appears in the frame. This design is safer for an academic project because it avoids overclaiming sensitive behavioural analysis while still adding a useful monitoring signal.")
subheading("B. Practical Use")
p("In a practical placement or internship screening scenario, the system can be used before a final technical interview. The administrator can prepare role-wise questions, share the test link, allow candidates to complete the assessment and then review generated reports. Candidates with stronger scores and acceptable monitoring status can be shortlisted for a human interview. This supports evaluators without removing human decision-making authority.")
subheading("C. Limitations")
p("The project has a few limitations. The descriptive scoring quality depends on the prepared ideal answer and keyword list. The system may not fully understand creative but correct answers if they are very different from the expected text. Camera monitoring depends on lighting, camera position and browser permission. Also, the current system does not include advanced authentication such as government ID verification or plagiarism checking. These limitations can be addressed in future versions.")

heading("VI. Conclusion and Future Work")
p("This paper presented an AI-Based Smart Interview System for automated first-level candidate screening. The system integrates Flask, PostgreSQL, NLP scoring, OpenCV monitoring and PDF report generation into a single workflow. The proposed approach supports role-based assessment, transparent scoring and structured administrator review. It is suitable for academic and internship demonstration and can be extended for organizational screening support.")
p("Future work includes adding HR login, improving semantic evaluation using BERT or sentence-transformer models, expanding the question bank, adding analytics dashboards, improving fairness checks and deploying the application on a secure server with stronger privacy controls.")
p("The system can also be extended with candidate analytics such as role-wise average score, topic-wise weakness, repeated question usage and interview completion trends. A more mature version can include consent screens, privacy policy enforcement, access control, encrypted report storage and review dashboards. These enhancements would make the system more suitable for real-world pilot deployment.")

heading("References")
refs = [
    "I. Naim et al., “Automated prediction and analysis of job interview performance,” 2015.",
    "L. Nguyen and D. Gatica-Perez, “Hirability in the wild: Analysis of online video resumes,” 2016.",
    "L. Chen et al., “Automated video interview judgment on large corpus,” 2017.",
    "S. Muralidhar et al., “Words worth: Verbal content and hirability impressions,” 2018.",
    "F. Pedregosa et al., “Scikit-learn: Machine learning in Python,” Journal of Machine Learning Research, 2011.",
    "K. Sparck Jones, “A statistical interpretation of term specificity and its application in retrieval,” 1972.",
    "G. Bradski, “The OpenCV Library,” Dr. Dobb's Journal of Software Tools, 2000.",
    "J. Devlin et al., “BERT: Pre-training of deep bidirectional transformers for language understanding,” 2019.",
    "H. Suen et al., “Artificial intelligence in personnel selection,” 2019.",
    "N. Guenole et al., “Algorithmic recruitment and selection systems,” 2022.",
]
for i, ref in enumerate(refs, 1):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.left_indent = Inches(0.18)
    para.paragraph_format.first_line_indent = Inches(-0.18)
    para.paragraph_format.space_after = Pt(1)
    para.paragraph_format.line_spacing = 1.0
    run = para.add_run(f"[{i}] {ref}")
    set_font(run, 8)

doc.save(OUT)
print(OUT)
