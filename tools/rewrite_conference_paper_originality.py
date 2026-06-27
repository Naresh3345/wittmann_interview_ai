from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


SOURCE = Path(r"E:\AI-Based Smart Interview System IEEE Paper.docx")
DOCX_OUT = Path(r"E:\AI-Based Smart Interview System IEEE Paper - Originality Revised.docx")


def replace_paragraph(paragraph, text):
    """Replace visible paragraph text while preserving the paragraph/run shell."""
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    if not tr_pr.xpath("./w:cantSplit"):
        cant_split = OxmlElement("w:cantSplit")
        cant_split.set(qn("w:val"), "1")
        tr_pr.append(cant_split)


PARAGRAPH_REWRITES = {
    0: "Explainable AI-Based Smart Interview System for First-Round Candidate Screening with NLP, Computer Vision and PostgreSQL",
    7: (
        "Abstract - Large recruitment drives and campus placement activities often require many similar first-round interviews. "
        "When every response is checked manually, evaluators spend considerable time and the feedback quality may vary from one candidate to another. "
        "This paper describes an AI-Based Smart Interview System built as a Flask web application for conducting an initial screening round. "
        "The system combines OTP-based candidate verification, role-wise question allocation, objective-answer validation, NLP-assisted descriptive answer scoring, "
        "OpenCV-based face-presence checks, PostgreSQL-backed record keeping and automatic report creation. Descriptive answers are assessed through keyword coverage, "
        "TF-IDF cosine similarity and a communication-quality component, so the result remains understandable to the administrator. "
        "The work focuses on a low-cost, locally deployable and auditable screening workflow that can support academic placement demonstrations and internship shortlisting without replacing the final human interview."
    ),
    8: "Keywords - Smart interview system; automated screening; explainable NLP; TF-IDF similarity; OpenCV monitoring; Flask application; PostgreSQL; interview report generation.",
    10: (
        "Initial interview screening helps organizations and placement teams identify applicants who show the required technical understanding, reasoning ability, communication clarity and professional readiness. "
        "In a traditional workflow, interviewers ask comparable questions to many candidates, listen to every response, compare answers with expectations and then prepare notes. "
        "Although this approach allows direct human judgement, it becomes slow and inconsistent when the same entry-level profile must be evaluated for a large candidate pool."
    ),
    11: (
        "Many online assessment platforms reduce manual effort only for objective questions. They are less useful when the evaluator wants to check descriptive technical answers, explain why a score was assigned or preserve interview evidence for later review. "
        "Video interview tools provide richer records, but they commonly depend on manual playback review or external paid analytics. "
        "A practical alternative is therefore needed: a transparent system that can run locally, manage the interview flow on the web, score answers in an explainable manner, observe simple camera-status signals, store records and produce reports."
    ),
    12: (
        "The proposed system addresses this need through an integrated web-based interview workflow. It includes candidate registration, OTP verification, role selection, role-specific question delivery, NLP-based scoring, OpenCV-based face-presence monitoring and PDF report generation. "
        "The current implementation supports roles such as Manual Testing, Automation Testing, AI Technical Support and Software Development, making it suitable for academic evaluation, internship screening and project demonstration."
    ),
    13: (
        "The central contribution of this work is the combination of several simple, interpretable components into one complete screening framework. "
        "Rather than treating the process as only a quiz portal or only a recorded video interview, the system links verification, question assignment, response evaluation, camera-state logging, database persistence and reporting in a single auditable pipeline."
    ),
    14: (
        "A second contribution is the emphasis on explainability. First-round screening results should be reviewable by a human evaluator; therefore, the report presents matched terms, missing terms, semantic similarity and communication score instead of only displaying a black-box decision. "
        "This design allows administrators to understand the basis of the score before using it for shortlisting."
    ),
    16: (
        "Previous research on automated interview analysis has explored verbal responses, speech patterns, facial cues and video behaviour. "
        "Studies such as those by Naim et al., Nguyen and Gatica-Perez, Chen et al. and Muralidhar et al. demonstrate that computational features can assist interview-performance estimation. "
        "However, many such approaches require large datasets, specialized feature extraction and complex models that are difficult to reproduce in a lightweight academic project."
    ),
    17: (
        "Automated answer assessment is also common in educational technology and recruitment support systems. Keyword matching and TF-IDF similarity are useful because their outputs can be explained and checked by the evaluator. "
        "Scikit-learn offers practical vectorization and similarity tools, while OpenCV supports real-time frame analysis tasks such as face detection. "
        "At the same time, literature on AI recruitment stresses transparency, fairness and human oversight. The proposed work follows that direction by using classical NLP and limited camera monitoring instead of opaque personality or emotion prediction."
    ),
    18: (
        "Existing tools may be grouped into manual interviews, online test portals and AI-enabled video interview platforms. Manual interviews provide rich judgement but consume evaluator time. Test portals scale well but usually focus on multiple-choice scoring. "
        "Video-interview platforms collect richer evidence but can be costly and review-heavy. The proposed system occupies a middle position by offering a locally deployable, explainable and role-based screening process."
    ),
    21: (
        "The system is structured as a layered web application. The presentation layer contains HTML, CSS and JavaScript interfaces for registration, OTP entry, role selection, question display, camera capture and result viewing. "
        "The application layer is implemented with Flask routes and API endpoints. The evaluation layer performs answer scoring and camera-frame analysis. The persistence and reporting layer stores information in PostgreSQL and creates PDF reports through ReportLab."
    ),
    25: (
        "A candidate begins by submitting basic identity and contact information along with resume details. The application generates a one-time password and allows the candidate to proceed only after successful verification. "
        "The selected role then determines the question mix, topics and evaluation pattern used for the interview session."
    ),
    27: (
        "The question bank is maintained in PostgreSQL. Each question record stores the role slug, section, topic, difficulty level, question statement, answer options, expected answer, keywords, marks and active status. "
        "When an interview starts, the application selects active questions for the chosen role and stores the assigned set for that session, so later question-bank changes do not alter completed interview records."
    ),
    29: (
        "Objective responses are checked through direct answer matching. Descriptive responses are scored using three signals: keyword coverage, TF-IDF cosine similarity against the expected answer and a communication score. "
        "The combined descriptive score is computed using the following weighted expression:"
    ),
    31: (
        "In this expression, KS is the keyword score, TS is the TF-IDF similarity score and CS is the communication score. "
        "The weighting gives semantic relevance the largest share while still recognizing important technical terms and the clarity of the candidate's response."
    ),
    33: (
        "During the interview, the browser captures webcam frames and sends encoded images to the Flask API. OpenCV converts each frame to grayscale and applies Haar-cascade face detection. "
        "The system records whether no face is visible, whether more than one face is detected and whether the candidate is generally centered. These values are used only as basic monitoring indicators."
    ),
    37: (
        "After submission, answer scores and camera-status summaries are combined into a candidate-level result. The generated report includes identity details, selected role, total score, question-level performance, feedback, confidence index and shortlist status. "
        "Administrators can access the stored report from the report page, which reduces the effort required to prepare interview notes manually."
    ),
    41: (
        "The backend is implemented in Python with Flask. The main application defines routes and APIs for the home page, OTP verification, role selection, interview page, database view, administrator reports, question retrieval, frame analysis, response submission, result display and report download. "
        "Supporting utility modules handle scoring, database operations, question selection and report generation."
    ),
    42: (
        "The frontend is built with HTML templates, CSS and JavaScript. The interview page loads assigned questions, stores answer selections, shows camera status, supports speech input from the browser and submits the completed attempt through an API call. "
        "The scoring utility uses scikit-learn's TF-IDF vectorizer and cosine similarity, the database utility connects to PostgreSQL through psycopg, and ReportLab creates the final PDF report."
    ),
    43: (
        "The codebase is organized so that core responsibilities remain separated. Scoring functions are kept in a utility module, allowing the evaluation logic to be tested and improved independently. "
        "Database functions create and maintain records for roles, users, interviews, questions and candidate answers. The question-bank module selects the least-used active questions for the chosen role. This separation improves maintainability and simplifies future extension."
    ),
    45: (
        "Auditability is an important part of the database design. The live question bank can be edited by an administrator, but the exact questions assigned to a candidate are copied into the interview record. "
        "As a result, completed results remain traceable even if a question is later modified, disabled or replaced."
    ),
    48: (
        "The system was validated through sample role-based candidate flows. The tested workflow included registration, OTP verification, role selection, question loading, answer submission, NLP scoring, camera monitoring, database storage, PDF report creation and administrator report viewing. "
        "At the reviewed stage, the implementation demonstrated roughly eighty percent completion of the planned project workflow."
    ),
    49: (
        "Test attempts were created for multiple roles. Each run checked that the correct question set was loaded, submitted answers were accepted by the API, score summaries were produced and report paths were stored correctly. "
        "Camera analysis was verified for normal face presence, missing-face cases and multiple-face alerts. These checks confirmed that the major modules can operate together as an end-to-end screening process."
    ),
    51: (
        "A major benefit of the system is its consistent and explainable screening process. PostgreSQL-based question management keeps the bank reusable and editable, while automatic reporting reduces documentation work. "
        "The system is not intended to replace final interviewer judgement. The current vision component is limited to simple monitoring and the NLP model uses classical similarity methods; future versions can improve semantic scoring, analytics and deployment security."
    ),
    53: (
        "The observed behaviour indicates that classical NLP can be useful for first-level screening. Exact matching is suitable for objective aptitude and programming questions. Keyword matching verifies whether important technical concepts are mentioned, while TF-IDF cosine similarity provides an approximate comparison between the submitted response and the expected answer. "
        "The communication component rewards responses that are reasonably clear and sufficiently developed."
    ),
    54: (
        "The OpenCV feature is deliberately conservative. It does not attempt to infer personality, emotion or intent. Instead, it reports only whether a candidate appears in the frame and whether multiple faces are present. "
        "This keeps the academic prototype from making sensitive behavioural claims while still adding a useful interview-monitoring signal."
    ),
    56: (
        "In a placement or internship-screening scenario, administrators can prepare role-wise questions, share the assessment link, allow candidates to complete the interview and then review the generated reports. "
        "Candidates with stronger scores and acceptable monitoring status can be moved to a human technical interview. In this way, the system supports evaluators while preserving human decision-making authority."
    ),
    58: (
        "The project has several limitations. Descriptive-score quality depends on the expected answer and keyword list prepared for each question. A correct but unusually phrased answer may receive a lower score if it differs greatly from the reference text. "
        "Camera monitoring is affected by lighting, camera placement and browser permissions. The current prototype also does not include government-ID verification, advanced proctoring or plagiarism checking. These gaps can be addressed in later versions."
    ),
    60: (
        "This paper presented an explainable AI-Based Smart Interview System for automated first-round candidate screening. "
        "The system integrates Flask, PostgreSQL, NLP scoring, OpenCV monitoring and PDF report generation into a single workflow. "
        "It supports role-based assessment, transparent scoring and structured administrator review, making it appropriate for academic demonstrations and internship-screening support."
    ),
    61: (
        "Future development can add an HR login module, improve semantic evaluation with BERT or sentence-transformer models, expand the role-wise question bank, introduce analytics dashboards, strengthen fairness checks and deploy the application on a secure server with stronger privacy controls."
    ),
    62: (
        "The system can also be extended with analytics such as role-wise average score, topic-wise weakness identification, question-usage trends and interview-completion statistics. "
        "A production-ready version should include consent screens, privacy-policy enforcement, role-based access control, encrypted report storage and richer review dashboards."
    ),
}


TABLE_REWRITES = {
    0: [
        ["Approach", "Primary Benefit", "Main Constraint"],
        ["Manual interview", "Rich evaluator judgement", "Difficult to scale for large pools"],
        ["Online test portal", "Quick objective-question evaluation", "Limited support for descriptive responses"],
        ["Video interview platform", "Preserves visual and spoken evidence", "Often costly and review-intensive"],
        ["Proposed system", "Combines NLP scoring, camera checks, database records and reports", "Uses basic AI; final judgement remains human"],
    ],
    1: [
        ["Step", "Operation"],
        ["1", "Register candidate and complete OTP verification"],
        ["2", "Load active PostgreSQL questions for the selected role"],
        ["3", "Persist the assigned question set for the session"],
        ["4", "Collect answers and camera-frame samples"],
        ["5", "Score exact, keyword, TF-IDF and communication signals"],
        ["6", "Summarize face-presence indicators"],
        ["7", "Create result, shortlist status and PDF report"],
    ],
    2: [
        ["Layer", "Main Elements", "Role in the System"],
        ["Frontend", "HTML, CSS, JavaScript", "Candidate input, question display, camera capture and submission"],
        ["Backend", "Flask routes and APIs", "Session control, OTP handling, validation and workflow coordination"],
        ["AI Evaluation", "Scikit-learn, OpenCV", "Answer scoring, similarity calculation, face detection and alerts"],
        ["Database", "PostgreSQL", "Roles, users, question bank, assigned papers and answer records"],
        ["Report", "ReportLab", "Score summary, feedback, confidence index and PDF output"],
    ],
    3: [
        ["Data Entity", "Purpose"],
        ["users", "Identity, contact details, OTP status and resume path"],
        ["roles", "Supported interview profiles and display names"],
        ["question_bank", "Active questions with topics, answers, keywords and marks"],
        ["interviews", "Tracks interview id, candidate id, chosen role, status and report path"],
        ["answers", "Responses, AI score, feedback and keyword details"],
    ],
    4: [
        ["Feature", "Observed Result"],
        ["OTP verification", "Candidate flow continues only after successful verification"],
        ["Role-wise questions", "Questions are retrieved from the PostgreSQL question bank"],
        ["NLP evaluation", "Exact, keyword, TF-IDF and communication scores are produced"],
        ["Camera monitoring", "Missing-face and multiple-face conditions are identified"],
        ["Report generation", "A PDF report is generated for administrator review"],
    ],
}


def main():
    doc = Document(SOURCE)

    for index, text in PARAGRAPH_REWRITES.items():
        replace_paragraph(doc.paragraphs[index], text)

    for table_index, rows in TABLE_REWRITES.items():
        table = doc.tables[table_index]
        for table_row in table.rows:
            prevent_row_split(table_row)
        for row_index, row_values in enumerate(rows):
            for col_index, value in enumerate(row_values):
                cell = table.rows[row_index].cells[col_index]
                replace_paragraph(cell.paragraphs[0], value)

    # Keep figure captions and IEEE headings centered after replacement.
    for idx in (0, 23, 35):
        doc.paragraphs[idx].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(DOCX_OUT)
    print(DOCX_OUT)


if __name__ == "__main__":
    main()
