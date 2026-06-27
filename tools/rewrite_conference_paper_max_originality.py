from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


SOURCE = Path(r"E:\AI-Based Smart Interview System IEEE Paper - Originality Revised.docx")
DOCX_OUT = Path(r"E:\AI-Based Smart Interview System IEEE Paper - Maximum Originality.docx")


def replace_paragraph(paragraph, text):
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    if not tr_pr.xpath("./w:cantSplit"):
        el = OxmlElement("w:cantSplit")
        el.set(qn("w:val"), "1")
        tr_pr.append(el)


TEXT = {
    0: "Locally Deployable Smart Interview Screening Assistant Using Flask, NLP, OpenCV and PostgreSQL",
    7: (
        "Abstract - First-stage interview screening is difficult to manage when many applicants are assessed for similar technical roles. "
        "A fully manual process requires repeated evaluation work, and the quality of comments can differ across reviewers. "
        "This paper presents a locally deployable smart interview assistant that conducts a structured screening round through a Flask-based web interface. "
        "The application verifies candidates with an OTP flow, maps each candidate to a selected role, serves questions from a PostgreSQL question bank, checks objective responses, evaluates descriptive answers with transparent NLP signals, records basic camera-presence indicators using OpenCV and prepares a report for administrator review. "
        "The descriptive-answer model combines term coverage, TF-IDF cosine comparison and response-clarity scoring, which makes the result easier to audit than a hidden prediction score. "
        "The system is intended for academic placement, internship review and demonstration use, where automation should reduce repetitive work while final selection remains a human decision."
    ),
    8: "Keywords - Interview automation; candidate assessment; transparent scoring; TF-IDF; OpenCV; Flask; PostgreSQL; placement screening.",
    10: (
        "Screening interviews are used to identify candidates who have the minimum technical knowledge, reasoning ability and communication quality required for the next selection stage. "
        "In many institutions and small recruitment teams, the same type of interview must be repeated for a large number of applicants. "
        "This repetition creates two practical problems: interviewers spend time on routine checking, and the judgement notes may not remain uniform across candidates."
    ),
    11: (
        "Conventional online tests solve only part of this problem because they usually focus on objective answers. "
        "They do not always handle descriptive technical explanations, show how a score was produced or preserve a complete assessment trail. "
        "Recorded video interviews add more evidence, but they still require review effort and may depend on external services. "
        "A lightweight system that joins web-based interviewing, answer analysis, basic camera-status checking, persistent storage and report generation can therefore be useful."
    ),
    12: (
        "The proposed application follows that approach. A candidate registers, verifies the OTP, chooses an interview role and completes the assigned question set. "
        "The backend then scores the responses, summarizes camera observations and creates a report. "
        "The prototype currently covers roles including Manual Testing, Automation Testing, AI Technical Support and Software Development, so it can be used in a controlled academic or internship-screening setting."
    ),
    13: (
        "The main value of the work is not a single complex model; it is the end-to-end connection of practical modules. "
        "Candidate verification, question selection, answer scoring, camera-state logging, database recording and PDF reporting are handled as one workflow. "
        "This makes the screening result easier to reproduce and review."
    ),
    14: (
        "The system also gives importance to score transparency. Instead of presenting only a pass/fail label, the report includes matched keywords, missing keywords, similarity value and communication score. "
        "These details help an administrator decide whether the automated result is reasonable before shortlisting the candidate."
    ),
    16: (
        "Research on automated interview evaluation has considered spoken content, facial information, language features and video behaviour. "
        "Prior studies by Naim et al., Nguyen and Gatica-Perez, Chen et al. and Muralidhar et al. show that computational features can support interview-performance analysis. "
        "Those methods, however, often need larger datasets and more advanced modelling than a small locally run academic prototype can reasonably maintain."
    ),
    17: (
        "Text-based answer evaluation is widely used in learning and assessment tools. Keyword checks and TF-IDF similarity are simple, explainable and practical for comparing a candidate response with a prepared reference answer. "
        "Scikit-learn provides the required vectorization and similarity functions, and OpenCV can perform basic frame-level face detection. "
        "Because AI hiring tools can raise fairness and transparency concerns, this project avoids sensitive personality or emotion claims and keeps the automated decision support limited and reviewable."
    ),
    18: (
        "Available recruitment-support solutions generally fall into three groups: human interviews, online test systems and video-interview platforms. "
        "Human interviews are rich but time-consuming, online tests are scalable but narrow, and video platforms preserve evidence but may be expensive or review-heavy. "
        "The proposed system is positioned between these options by offering a low-cost, local and explainable first-round workflow."
    ),
    21: (
        "The application is divided into four working layers. The browser layer handles registration, OTP entry, role selection, question display, answer entry, camera capture and result viewing. "
        "The Flask layer manages routes, sessions, validation and API calls. The evaluation layer calculates response scores and camera indicators. "
        "The storage and reporting layer keeps records in PostgreSQL and generates administrator-facing PDF output using ReportLab."
    ),
    25: (
        "The candidate first enters personal and contact details. After OTP verification, the candidate is allowed to choose the interview role. "
        "That role controls which topics and question types are assigned during the attempt."
    ),
    27: (
        "Questions are stored as PostgreSQL records with role, section, topic, difficulty, question text, options, expected answer, keywords, marks and activity status. "
        "At the start of an interview, active questions are selected and copied into the candidate's interview record. "
        "This copy protects completed attempts from later edits made to the live question bank."
    ),
    29: (
        "The scoring module treats objective and descriptive answers differently. Objective responses are verified by answer matching. "
        "For descriptive responses, the system calculates keyword coverage, TF-IDF cosine similarity and communication quality, and then combines them with a fixed weight:"
    ),
    31: (
        "Here, KS represents keyword coverage, TS represents text-similarity score and CS represents communication score. "
        "The formula gives more influence to semantic closeness while still rewarding important technical terms and readable expression."
    ),
    33: (
        "Camera monitoring is implemented as a simple status check. Frames captured in the browser are sent to the backend, converted to grayscale and processed with Haar-cascade face detection. "
        "The stored indicators show whether a face was missing, whether multiple faces appeared and whether the candidate stayed generally visible. "
        "These signals are used only as basic confidence information."
    ),
    37: (
        "When the candidate submits the attempt, the application combines answer-level scores with the camera summary. "
        "The generated report lists candidate information, selected role, total marks, question-wise feedback, confidence index and shortlist status. "
        "Because the report is stored and accessible from the administrator page, evaluators do not need to prepare the same summary manually."
    ),
    41: (
        "The backend is written in Python using Flask. It includes endpoints and routes for landing, OTP verification, role selection, interview display, database inspection, reports, question loading, frame analysis, answer submission, result viewing and report download. "
        "Separate helper modules manage score calculation, database access, question selection and PDF creation."
    ),
    42: (
        "The client side uses HTML templates with CSS and JavaScript. During an interview, JavaScript loads the assigned questions, records answer choices, displays camera status, supports browser speech input and sends the completed response set to the backend. "
        "NLP scoring is implemented with scikit-learn, PostgreSQL connectivity is handled through psycopg and final reports are prepared with ReportLab."
    ),
    43: (
        "The implementation is organized around maintainability. Scoring logic is kept away from routing code, database operations are grouped in utility functions, and question selection is handled by a dedicated module. "
        "This layout makes it easier to update a scoring rule, extend the question bank or add a new role without rewriting the full application."
    ),
    45: (
        "The database design supports later review. Although administrators can update the active question bank, the questions assigned to a candidate are preserved inside that interview record. "
        "Therefore, an old report can still be traced to the exact questions that were actually answered."
    ),
    48: (
        "The prototype was checked with sample candidate attempts across different roles. The checked flow covered registration, OTP validation, role choice, question delivery, answer submission, NLP scoring, camera checks, database updates, report creation and administrator viewing. "
        "At the review stage, the project had implemented most of the planned screening workflow."
    ),
    49: (
        "During testing, each attempt was examined for correct question loading, successful API submission, generation of score summaries and proper report-path storage. "
        "The camera module was also tested for ordinary face visibility, no-face conditions and multiple-face alerts. "
        "These checks showed that the main modules work together as a complete screening path."
    ),
    51: (
        "The key advantage of the system is a repeatable screening method with understandable scoring evidence. "
        "PostgreSQL makes the question bank reusable, and report generation reduces documentation effort. "
        "The system is still a decision-support tool rather than a replacement for interviewers. "
        "The current NLP and vision methods are intentionally basic and can be improved in future versions."
    ),
    53: (
        "The results suggest that classical NLP is sufficient for a controlled first-round filter. Exact matching works for aptitude and programming answers, keyword checks confirm the presence of important concepts, and TF-IDF similarity provides a rough comparison with the reference answer. "
        "The communication score adds a small adjustment for clarity and completeness."
    ),
    54: (
        "The computer-vision part is kept narrow by design. It does not classify emotion, attention or personality. "
        "It only reports face presence and possible multiple-person cases. "
        "This limited use is more appropriate for an academic prototype and avoids unsupported behavioural conclusions."
    ),
    56: (
        "In a practical placement workflow, an administrator can configure questions, distribute the assessment link, collect completed attempts and review the generated reports. "
        "Candidates with acceptable scores and monitoring status can then be called for a human technical round. "
        "The system therefore reduces preliminary workload while keeping final selection under human control."
    ),
    58: (
        "The prototype has limitations. Descriptive scoring depends strongly on the quality of the prepared reference answer and keyword list. "
        "A valid answer written in an unusual style may not receive full credit. Camera results can be affected by lighting, device position and browser permission. "
        "The present version also does not include government identity checks, advanced proctoring or plagiarism detection."
    ),
    60: (
        "This paper described a locally deployable smart interview screening assistant that combines Flask, PostgreSQL, transparent NLP scoring, OpenCV-based camera checks and PDF reporting. "
        "The system supports role-based first-round assessment and gives administrators reviewable score evidence. "
        "It is suitable for academic project demonstration and controlled internship-screening support."
    ),
    61: (
        "Future improvements may include an HR login module, stronger semantic scoring with transformer embeddings, a larger role-wise question bank, dashboard analytics, fairness review features and secure server deployment with privacy controls."
    ),
    62: (
        "Additional extensions can include role-level score trends, topic-wise weakness reports, question-usage analysis and interview-completion metrics. "
        "A production version should also add consent handling, access control, encrypted report storage and richer administrator dashboards."
    ),
}


TABLES = {
    0: [
        ["Method", "Useful Point", "Weak Point"],
        ["Human interview", "Detailed judgement", "Slow for large groups"],
        ["Online quiz", "Fast checking", "Mostly objective answers"],
        ["Video platform", "Keeps richer evidence", "Cost and review effort"],
        ["This work", "Local NLP, camera status, database and report flow", "Basic models; human review needed"],
    ],
    1: [
        ["Step", "Task"],
        ["1", "Verify candidate through OTP"],
        ["2", "Load role-based active questions"],
        ["3", "Save assigned question set"],
        ["4", "Collect answers and camera frames"],
        ["5", "Compute answer-score components"],
        ["6", "Record camera-status indicators"],
        ["7", "Prepare shortlist result and report"],
    ],
    2: [
        ["Layer", "Elements", "Function"],
        ["Client", "HTML, CSS, JavaScript", "Forms, questions, answers and camera capture"],
        ["Server", "Flask APIs", "Sessions, validation and workflow control"],
        ["Evaluation", "Scikit-learn, OpenCV", "Text scoring and face-status checks"],
        ["Storage", "PostgreSQL", "Users, roles, questions, interviews and answers"],
        ["Output", "ReportLab", "Feedback, confidence index and PDF report"],
    ],
    3: [
        ["Entity", "Use"],
        ["users", "Candidate details and verification status"],
        ["roles", "Interview profile names"],
        ["question_bank", "Active questions, answers, keywords and marks"],
        ["interviews", "Attempt, role, status and report path"],
        ["answers", "Responses, scores and feedback"],
    ],
    4: [
        ["Module", "Result"],
        ["OTP", "Only verified candidates continue"],
        ["Questions", "Role-based items load from PostgreSQL"],
        ["Scoring", "Exact, keyword, similarity and clarity scores are produced"],
        ["Camera", "Missing-face and multiple-face states are flagged"],
        ["Report", "Administrator PDF output is created"],
    ],
}


def main():
    doc = Document(SOURCE)

    for index, text in TEXT.items():
        replace_paragraph(doc.paragraphs[index], text)

    for table_index, rows in TABLES.items():
        table = doc.tables[table_index]
        for table_row in table.rows:
            prevent_row_split(table_row)
        for r, values in enumerate(rows):
            for c, value in enumerate(values):
                replace_paragraph(table.rows[r].cells[c].paragraphs[0], value)

    for idx in (0, 23, 35):
        doc.paragraphs[idx].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(DOCX_OUT)
    print(DOCX_OUT)


if __name__ == "__main__":
    main()
